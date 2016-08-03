import xlrd
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

document = Document()
activities_template_workbook = xlrd.open_workbook('../U1L01 Final Template.xlsx', encoding_override="cp1252")

def create_table(data_list):
	table = document.add_table(rows=len(data_list), cols=1)
	for each in data_list:
		rv = data_list.index(each)
		# print "\n\n rv = ", rv
		cell = table.cell(0, rv)
		cell.text = "Test labels" + str(rv)
		# p = cell.paragraphs[0]
		# run = p.add_run("Test labels")
		# run.bold = True
		# run.font.color.rgb = RGBColor(0xff, 0, 0)
		# run.font.size = Pt(10)
		# run.font.name = 'Arial'
		# Set a cell background (shading) color to RGB D9D9D9. 
		shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
		cell._tc.get_or_add_tcPr().append(shading_elm)

document.add_heading('Document Title', level=5)
data_set = ['sahfgd', 'mdhgsfjd']
create_table(data_set)
document.save('testmatser1.docx')
print "\n doc saved"